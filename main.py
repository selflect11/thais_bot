import re
import code
from typing import List
from config import *
from itertools import tee
from docx import Document


def pairwise(iterable):
    "s -> (s0, s1), (s1, s2), (s2, s3), ..."
    a, b = tee(iterable)
    next(b, None)
    return zip(a, b)

def get_page_blocks(paragraph_list):
    # receives paragraph objects rather than text

    def get_page_indexes(paragraph_list):
        header_pat = re.compile(MAIN_HEADER)
        header_indexes = []
        for i, paragraph in enumerate(paragraph_list):
            if re.search(header_pat, paragraph.text):
                header_indexes.append(i)
        return header_indexes

    pages = []
    header_indexes = get_page_indexes(paragraph_list)
    for i_0, i_1 in pairwise(header_indexes):
        pages.append(paragraph_list[i_0:i_1])
    # last page doesnt have last+1 corresponding header index, so we do this
    last_i = header_indexes[-1]
    pages.append(paragraph_list[last_i:-1])
    return pages

def generate_header(page):
    ano, banca, orgao = '', '', ''
    for paragraph in page:
        search_banca = re.search(BANCA_HEADER, paragraph.text)
        search_orgao = re.search(ORGAO_HEADER, paragraph.text)
        search_ano = re.search(ANO_HEADER, paragraph.text)
        if search_banca:
            banca = search_banca.group('banca')
            # orgao not always present. When it is, it's on the same paragraph as banca
            if search_orgao:
                orgao = search_orgao.group('orgao')
        elif search_ano:
            ano = search_ano.group('ano')
        elif ano and banca and orgao:
            break
    if orgao:
        return '({orgao} - {banca} - {ano}) '.format(ano=ano, banca=banca, orgao=orgao)
    else:
        return '({banca} - {ano}) '.format(ano=ano, banca=banca)

def field_filter(page):
    # returns only fields that are supposed to be inaltered
    empty_pat = re.compile(r'(?:\s+)|(?:^$)')
    between_disc_and_nivel = False
    filtered_page = []

    for paragraph in page:
        is_empty = re.match(empty_pat, paragraph.text)
        delete_paragraphs = [re.search(pat, paragraph.text) for pat in DELETE_HEADERS]
        if re.search(DISCIPLINA_HEADER, paragraph.text):
            between_disc_and_nivel = True
        elif re.search(NIVEL_HEADER, paragraph.text):
            between_disc_and_nivel = False
        elif not any(delete_paragraphs) and not is_empty and not between_disc_and_nivel:
            filtered_page.append(paragraph)
    return filtered_page

def field_reorganize(page):
    # Move 'Dica do autor' to before 'Referencias bibliograficas'
    # 'Dica do autor' ends when 'Alternativa A:' starts
    organized_paragraphs = []

    def get_dicas(page):
        is_dicas_block = False
        dicas = []
        for paragraph in page:
            if re.search(DICA_HEADER, paragraph.text):
                is_dicas_block = True
                dicas.append(paragraph)
                break
            elif re.search(QUESTOES_HEADER, paragraph.text):
                # for readability
                is_dicas_block = False
                break
            elif is_dicas_block:
                dicas.append(paragraph)
        return dicas

    dicas = get_dicas(page)
    # moves dica to before 'referencias'
    for paragraph in page:
        if re.search(REFERENCIAS_HEADER, paragraph.text):
            for dica in dicas:
                organized_paragraphs.append(dica)
        elif paragraph.text not in dicas:
            organized_paragraphs.append(paragraph)
    return organized_paragraphs

def format_doc(page_list):
    num_fmt = '{num:02d}. '
    out_pages = []
    for i, page in enumerate(page_list):
        head = generate_header(page)
        filtered_page = field_filter(page)
        fixed_page = field_reorganize(filtered_page)
        if not fixed_page:
            break
        n = num_fmt.format(num=i+1)

        # 01. (orgao - banca - ano) ← start of first paragraph
        # 'Enunciado' is the rest
        header_paragraph = Document().add_paragraph(n + head + fixed_page[0].text)
        # out_page = [parag11, parag21, parag31, ...]
        out_page = [header_paragraph] + fixed_page[1:]
        # out_pages = [[parag11, parag21, parag31, ...], [parag12, parag22, parag32, ...] ...]
        out_pages.append(out_page)
    return out_pages

def save_doc(path, page_list):
    d = Document()
    for page in page_list:
        for paragraph in page:
            # d.add_paragraph(paragraph)
            # ugly ass hack
            d._body._body._insert_p(paragraph._p)
        # WARNING: documentation doesnt advise overuse of this. Dont know why. Investigate...
        # It seems like some pages already have this
        d.add_page_break()
    d.save(path)

if __name__ == '__main__':
    paf = FILE_INPUT_PATH
    d = Document(FILE_INPUT_PATH)
    pages = get_page_blocks(d.paragraphs)
    fixed_pages = format_doc(pages)
    save_doc(paf, fixed_pages)

    if DEBUG:
        code.interact(local=locals())
